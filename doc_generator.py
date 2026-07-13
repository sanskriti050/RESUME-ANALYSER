from docx import Document
from io import BytesIO


def create_docx(title, content):
    doc = Document()

    doc.add_heading(title, level=1)

    for line in content.split("\n"):
        doc.add_paragraph(line)

    file = BytesIO()
    doc.save(file)
    file.seek(0)

    return file